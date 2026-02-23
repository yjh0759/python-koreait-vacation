# 주독일대사관_크롤링
import requests
from bs4 import BeautifulSoup
from time import sleep
import os
import docx
from docx import Document

# docx 폰트지정
def apply_font(*texts):
    for text in texts:
        text.runs[0].font.name = "Malgun Gothic"


# 1. url패턴 파악
BASE_ID = 2975124
# ?seq=2975123&page=1 // ? 이후부분은 추가정보요청부분(생략 가능할 수 있음)
BASE_URL= "https://bonn.mofa.go.kr/de-ko/brd/m_7204/view.do?seq="
COUNT = 10
BASE_ROOT = "./독일대사관_게시글" # 저장할 폴더위치
# 2. url 조립
target_urls = []
for num in range(0, COUNT):
    url = f"{BASE_URL}{BASE_ID - num}"
    print(url)
    target_urls.append(url)

# 3. 게시글 크롤링
count = 0
for url in target_urls:
    response = requests.get(url)
    # 통신코드 -> 200이면 성공!
    print(response.status_code)
    html = BeautifulSoup(response.text, "html.parser")

    header_tag = html.find("div", class_="bo_head")


    # 에러페이지 처리
    if header_tag is None:
        print("    -> 에러페이지 발생!! 건너뜁니다")
        count += 1
        continue

    # 게시글 제목 추출
    title = header_tag.find("h2").text.strip()
    print(f"제목: {title}")


    # 게시글 작성자, 작성일자 추출
    # dd태그라는 것 외에 식별가능한게 "순서"
    dd_tags = header_tag.find_all("dd")
    print(len(dd_tags)) # 2
    # 작성자
    author = dd_tags[0].text.strip()
    # 작성일자
    date = dd_tags[1].text.strip()
    print(f"작성자: {author}")
    print(f"작성일자: {date}")

    # 도전! - 본문내용 추출
    # 힌트) 상위태그의 text에는 하위태그의 text들도 모두 포함되어있음!
    content= html.find("div", class_="bo_con").text.strip()
    print(f"내용:{content[:100]}...")

    ##### docx 파일로 저장 #####
    doc = Document()
    doc.add_heading(title, level=1)
    doc.add_paragraph(f"작성일자: {date}")
    doc.add_paragraph("")
    pg3 = doc.add_paragraph(content)


    # 폴더생성
    os.makedirs(BASE_ROOT, exist_ok=True)
    # 파일저장 경로 조립
    file_path = f"{BASE_ROOT}/{title}_{date}.docx"

    doc.save(file_path)
    print(f"   -> 저장완료: {file_path}")
    sleep(1.5)

print("크롤링 완료")
success_count = COUNT - count
print(f"총 {success_count}건 크롤링!")







#https://bonn.mofa.go.kr/de-ko/brd/m_7204/view.do?seq=2975123&page=1
#https://bonn.mofa.go.kr/de-ko/brd/m_7204/view.do?seq=2975121&page=1
#https://bonn.mofa.go.kr/de-ko/brd/m_7204/view.do?seq=2975122&page=1