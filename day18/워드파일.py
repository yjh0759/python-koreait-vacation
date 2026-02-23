# 워드파일 파이썬으로 docx 생성
# pip install python-docx

from docx import Document

# 빈 워드문서 생성
doc = Document()

# 한글폰트가 깨질 수 있음(환경설정 문제)
# 폰트지정하는 함수
def apply_font(*texts):
    # *text: 여러 매개변수를 tuple로 감싸서 함수안으로 들고옴
    for text in texts:
        text.runs[0].font.name = "Malgun Gothic"

# 제목 추가 - level(제목크기 1,2,3)
title = doc.add_heading("파이썬스터디", level=1)

# 본문 추가 - 한줄씩 추가
pg1 = doc.add_paragraph("크롤링 공부중")
pg2 = doc.add_paragraph("폰트 깨는지 테스트중")

apply_font(title, pg1, pg2)

# docx 파일 저장
doc.save("./파이썬테스트.docx")


##### os 모듈 #####
# 파일경로, 폴더를 파이썬으로 다룰 때
import os

# 1. 파일이나 폴더가 경로에 존재하는가?
# os.path.exists(경로)
if os.path.exists("/파이썬테스트.docx"):
    print("파이썬테스트 파일이 존재합니다!")

if not os.path.exists("./자바테스트.docx"):
    print("자바테스트 파일은 존재하지 않습니다!")

# 2. 폴더 만들기
# makedirs()
# exist_ok = True : 이미 있어도 괜찮냐?
os.makedirs("내가만든폴더", exist_ok=True)

# 연습) docx파일을 생성해 주세요
# 제목: "내가만든 docx", 글: "모듈 연습중입니다"
# 내가만든폴더2에 연습.docx 이름으로 저장해 주세요
new_doc = Document()

title = new_doc.add_heading("내가만든 docx", level=1)
pg = new_doc.add_paragraph("모듈 연습중입니다")
apply_font(title, pg)

os.makedirs("내가만든폴더2", exist_ok=True)

if os.path.exists("./내가만든폴더2"):
    new_doc.save("./내가만든폴더2/연습.docx")











